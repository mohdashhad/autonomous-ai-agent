import os
from docx import Document
from datetime import datetime

def create_word_document(title: str, content: str) -> str:
    """
    Creates a Microsoft Word (.docx) document and saves it locally.
    Returns the generated file path.
    """
    doc = Document()
    
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    
    os.makedirs("output", exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = title.replace(' ', '_').replace('/', '_')[:20]
    filename = f"output/{safe_title}_{timestamp}.docx"
    
    doc.save(filename)
    return filename